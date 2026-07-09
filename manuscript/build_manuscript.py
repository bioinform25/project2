#!/usr/bin/env python3
"""Build the main manuscript docx from draft_content.md using python-docx.
No pandoc/LibreOffice available on this machine, so content is built directly
via the python-docx API rather than converted from markdown."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = r"C:\Users\SAMSUNG\Desktop\project2"
FIG = os.path.join(BASE, "figures")
OUT = os.path.join(BASE, "manuscript", "manuscript.docx")

doc = Document()

# ---- base style ----
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.15

for level, size in [(1, 15), (2, 13), (3, 12)]:
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Times New Roman"
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.space_before = Pt(14)
    hs.paragraph_format.space_after = Pt(6)


def add_caption(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.italic = True
    return p


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(headers, rows, col_widths=None, header_shade="D9D9D9", font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        set_cell_shading(hdr_cells[i], header_shade)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(path, width_in=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))


# =====================================================================
# TITLE PAGE
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run(
    "HSF2 Expression Is Inversely, Not Positively, Associated with a "
    "NETosis-Core Gene Signature Across Liver and Lung Fibrosis: "
    "A Multi-Cohort Transcriptomic Analysis"
)
r.bold = True
r.font.size = Pt(16)
title.paragraph_format.space_after = Pt(18)

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.add_run("[Author Name]").font.superscript = False
run = authors.add_run("[Author Name]1, [Author Name]1")
run.font.size = Pt(12)
authors.paragraph_format.space_after = Pt(4)

affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = affil.add_run("1[Department], [Institution]")
run.font.size = Pt(11)
run.italic = True
affil.paragraph_format.space_after = Pt(24)

corr = doc.add_paragraph()
corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = corr.add_run("Corresponding author: [Author Name], [Email]")
run.font.size = Pt(10)
corr.paragraph_format.space_after = Pt(24)

# =====================================================================
# ABSTRACT
# =====================================================================
doc.add_heading("Abstract", level=1)

abstract_parts = [
    ("Background: ", "Environmental-toxicant-exposed neutrophils have been proposed to adopt a "
     "pro-fibrotic, immunosuppressive “N2” phenotype characterized by neutrophil extracellular "
     "trap (NET) release, a process potentially relevant to hepatocellular carcinoma (HCC) immune "
     "evasion. In a pilot experiment, the stress/development-associated transcription factor heat "
     "shock factor 2 (HSF2) was found to increase in polyhexamethylene guanidine (PHMG)-treated "
     "neutrophils, motivating the hypothesis that HSF2 upregulation drives NETosis-associated gene "
     "expression during toxicant-induced fibrosis."),
    ("Methods: ", "Before further wet-laboratory work, we tested the population-level prediction of "
     "this hypothesis — that HSF2 should positively correlate with NETosis-core gene expression "
     "(PADI4, ELANE, MPO) in fibrotic tissue — using four independent public transcriptomic "
     "cohorts (two liver: GSE135251, GSE14323; two lung: GSE47460, GSE53845) spanning three "
     "microarray/RNA-seq platforms. Pre-specified Spearman correlation, bootstrap confidence "
     "intervals, Benjamini-Hochberg FDR correction, and random-effects meta-analysis "
     "(DerSimonian-Laird) were used. A PTPRC (CD45) specificity control, ChEA3 in silico "
     "transcription-factor enrichment, and an exploratory analysis of three independent HCC "
     "immune-checkpoint-inhibitor (ICI) cohorts were performed as follow-up."),
    ("Results: ", "Contrary to the original hypothesis, HSF2 showed a significant negative "
     "correlation with the NETosis-core score in 3 of 4 cohorts (pooled random-effects rho = "
     "−0.31, 95% CI [−0.51, −0.08], p = 0.010), which was not explained by generic "
     "immune-cell infiltration (PTPRC showed the opposite, positive association). No existing "
     "ChEA3 evidence library supported HSF2 as a direct positive regulator of these genes. Across "
     "three independent HCC-ICI cohorts, lower tumor HSF2 was nominally associated with treatment "
     "response in two combination-therapy cohorts, but pooled meta-analysis was not significant "
     "(p = 0.22) due to high heterogeneity."),
    ("Conclusions: ", "Public transcriptomic data do not support the hypothesis that HSF2 drives "
     "NETosis-associated gene expression in fibrotic tissue; instead, they suggest an unexplained "
     "inverse association warranting direct testing in isolated, toxicant-exposed neutrophils."),
]
p = doc.add_paragraph()
for bold, text in abstract_parts:
    rb = p.add_run(bold)
    rb.bold = True
    p.add_run(text + " ")

kw = doc.add_paragraph()
rk = kw.add_run("Keywords: ")
rk.bold = True
kw.add_run("HSF2; neutrophil extracellular traps; NETosis; N1/N2 neutrophil polarization; liver "
            "fibrosis; pulmonary fibrosis; hepatocellular carcinoma; immune checkpoint inhibitor; "
            "meta-analysis")

doc.add_page_break()

# =====================================================================
# 1. INTRODUCTION
# =====================================================================
doc.add_heading("1. Introduction", level=1)

intro_paragraphs = [
 "Hepatocellular carcinoma (HCC) remains one of the leading causes of cancer-related mortality "
 "worldwide, and although immune checkpoint inhibitors (ICIs) have become a standard component of "
 "systemic therapy for advanced disease, a substantial fraction of patients show primary or "
 "acquired resistance, motivating an ongoing search for predictive biomarkers and novel "
 "therapeutic targets (1). Neutrophils are increasingly recognized as functionally heterogeneous "
 "participants in this resistance: tumor-associated neutrophils can be polarized toward a "
 "cytotoxic, antitumor “N1” phenotype or a pro-tumorigenic, immunosuppressive “N2” "
 "phenotype, a dichotomy first described in the context of TGF-β signaling (2) and subsequently "
 "reproduced in vitro (3). N2-polarized and otherwise activated neutrophils can release neutrophil "
 "extracellular traps (NETs) — web-like structures of decondensed chromatin decorated with "
 "granule proteins such as myeloperoxidase (MPO) and neutrophil elastase (ELANE), a process that "
 "requires the citrullinating enzyme peptidylarginine deiminase 4 (PAD4/PADI4) (4). Excessive NET "
 "formation has been mechanistically linked to fibrosis progression and to tumor immune evasion, "
 "making the transcriptional control of NETosis-associated genes of direct translational interest.",

 "Polyhexamethylene guanidine (PHMG), an environmental disinfectant toxicant with a "
 "well-documented history of causing severe lung injury and fibrosis, induces "
 "endoplasmic-reticulum-stress-mediated epithelial apoptosis (5) and robust pulmonary neutrophil "
 "infiltration and fibrosis in animal models (6). In a pilot laboratory experiment, primary "
 "neutrophils treated with PHMG (3 μg/mL, 6 hours) showed a visibly increased protein band for "
 "heat shock factor 2 (HSF2) relative to untreated controls. HSF2 is a paralog of the canonical "
 "acute-stress-responsive transcription factor HSF1; unlike HSF1, HSF2 responds weakly to acute "
 "heat stress and is instead more classically associated with proteasome-inhibition-triggered "
 "stress, cellular differentiation, and development, where it can heterotrimerize with and "
 "fine-tune HSF1 activity (7, 8). More recently, HSF2 has been implicated in cancer biology in a "
 "context-dependent, sometimes opposing manner: it suppresses invasion in prostate cancer (9), "
 "while cooperating with HSF1 to sustain a pro-malignant transcriptional program across multiple "
 "cancer types (10), and a pan-cancer bioinformatic analysis found HSF2 to be an independent "
 "poor-prognosis factor in liver hepatocellular carcinoma (LIHC) whose expression correlates "
 "positively with immune-checkpoint genes (PD-1, PD-L1, CTLA4) and with several immune-cell "
 "infiltration signatures (11).",

 "Building on the pilot finding, we formulated the hypothesis that PHMG-induced HSF2 upregulation "
 "drives neutrophils toward an N2/NETotic phenotype, thereby promoting fibrosis and, in the liver, "
 "an immunosuppressive microenvironment permissive to HCC immune evasion. Because no existing "
 "literature directly connects HSF2 to neutrophil polarization or NETosis (confirmed by literature "
 "review prior to this study), and because ChIP-seq evidence for HSF2 binding is sparse and not "
 "derived from myeloid cells, this hypothesis is not yet supported by direct mechanistic evidence "
 "in any cell type. Before committing further wet-laboratory resources to this specific "
 "mechanistic model, we reasoned that the hypothesis makes a testable population-level prediction: "
 "if HSF2 upregulation drives NETosis-associated gene expression, tissue-level HSF2 expression "
 "should positively correlate with the expression of core NETosis genes (PADI4, ELANE, MPO) "
 "specifically in fibrotic tissue, and this relationship should be reproducible across independent "
 "organs and cohorts. We tested this prediction using publicly available transcriptomic data from "
 "human liver and lung fibrosis cohorts, applying a pre-specified statistical plan with formal "
 "replication criteria, and extended the analysis with a specificity control, an in silico "
 "transcription-factor-target enrichment check, and an exploratory analysis of HSF2 in the context "
 "of HCC immunotherapy response.",
]
for para in intro_paragraphs:
    doc.add_paragraph(para)

# =====================================================================
# 2. METHODS
# =====================================================================
doc.add_heading("2. Methods", level=1)

doc.add_heading("2.1 Pre-specified hypothesis and analysis plan", level=2)
doc.add_paragraph(
 "The analysis plan was fixed before inspecting any correlation result (documented in the "
 "project's version-controlled code repository). The primary endpoint was the Spearman "
 "correlation between HSF2 expression and a NETosis-core gene score (mean of gene-wise z-scores "
 "for PADI4, ELANE, and MPO — the three genes encoding proteins mechanistically required for "
 "NET formation) within diseased/fibrotic tissue samples in each cohort. A secondary "
 "“NETosis-extended” score added CTSG, LTF, CAMP, DEFA4, AZU1, ITGAM, NCF2, S100A12, and "
 "LCN2 (broader neutrophil granule/antimicrobial genes). Two independent cohorts per organ (liver, "
 "lung) were designated a priori as “primary” (larger n or continuous disease staging) and "
 "“replication” (independent platform). A finding was classified as replicated only if the "
 "primary and replication cohort agreed in sign and were each nominally significant (p < 0.05). "
 "Individual-gene correlations were treated as secondary/exploratory and corrected for multiple "
 "comparisons (Benjamini-Hochberg false discovery rate, FDR)."
)

doc.add_heading("2.2 Data sources", level=2)
doc.add_paragraph(
 "Four publicly available fibrosis cohorts were obtained from the NCBI Gene Expression Omnibus "
 "(GEO) (12) using the GEOquery Bioconductor package (13) (Table 1). Liver: GSE135251, a bulk "
 "RNA-seq cohort spanning the non-alcoholic fatty liver disease (NAFLD) severity spectrum with "
 "histological fibrosis staging F0–F4 (14), used as the primary liver cohort; and GSE14323, an "
 "Affymetrix HG-U133A/GPL571 microarray cohort of normal, cirrhotic, cirrhotic-with-HCC, and HCC "
 "liver tissue (15), used as the replication liver cohort. Lung: GSE47460, an Agilent (GPL14550 "
 "sub-series) cohort from the Lung Genomics Research Consortium/Lung Tissue Research Consortium "
 "comprising control, chronic obstructive pulmonary disease (COPD), and interstitial lung disease "
 "(ILD) subjects, of whom the histologically fibrotic “usual interstitial "
 "pneumonia/idiopathic pulmonary fibrosis” (UIP/IPF) subtype was used as the fibrotic "
 "population, used as the primary lung cohort (a single peer-reviewed publication describing this "
 "specific consortium sub-series could not be unambiguously identified among the several linked "
 "citations; the dataset is cited by GEO accession); and GSE53845, a two-color Agilent (GPL6480) "
 "cohort of IPF and control lung tissue (16), used as the replication lung cohort."
)
doc.add_paragraph(
 "Three independent HCC immune-checkpoint-inhibitor (ICI) response cohorts were analyzed as an "
 "exploratory extension: GSE215011 (nivolumab monotherapy, whole-transcriptome tumor RNA-seq, "
 "n = 10) (17); GSE279750 (anti-PD-L1-based combination therapy, post-treatment tumor RNA-seq, "
 "n = 10; no peer-reviewed publication was available in the GEO record at the time of access, "
 "cited by accession only) (23); and GSE235863 (anti-PD-1 plus lenvatinib combination therapy, "
 "post-treatment bulk tumor RNA-seq TPM matrix, n = 15 patients with available bulk data) (18). A "
 "fourth candidate ICI cohort, GSE140901 (19), was evaluated but excluded because its 785-gene "
 "NanoString PanCancer Immune Profiling panel does not include HSF2 and is missing two of the "
 "three NETosis-core genes (PADI4, MPO)."
)

doc.add_heading("2.3 Data processing", level=2)
doc.add_paragraph(
 "For RNA-seq count data (GSE135251), gene-level Ensembl counts were obtained per sample, genes "
 "with mean count < 10 across samples were removed except for pre-specified panel genes (which "
 "were retained regardless of this filter to accommodate the comparatively low bulk-tissue "
 "expression of neutrophil-lineage transcripts, with per-gene detection rates reported explicitly "
 "in Supplementary Table S1d), and variance-stabilizing transformation (blind to any grouping "
 "variable) was applied using DESeq2 (20). For microarray data (GSE14323, GSE47460, GSE53845), "
 "previously normalized expression values provided by data submitters were used as-is (confirmed "
 "log2-scale; RMA for GSE14323, cyclic-loess for GSE47460, two-color log-ratio for GSE53845). "
 "Where a gene symbol mapped to multiple Ensembl IDs or probes, the ID/probe with the highest mean "
 "expression across the cohort was retained. For the HCC-ICI cohorts (GSE215011, GSE279750: "
 "per-sample tables; GSE235863: a shared TPM matrix), gene-level expression was log2(x+1)-"
 "transformed."
)

doc.add_heading("2.4 Statistical analysis", level=2)
doc.add_paragraph(
 "Spearman correlation coefficients were computed for the primary and secondary endpoints, "
 "accompanied by 2000-replicate percentile bootstrap 95% confidence intervals and leave-one-out "
 "sensitivity analysis. A specificity control tested whether HSF2 also correlated with PTPRC "
 "(CD45, a pan-leukocyte marker) to distinguish a NETosis-specific transcriptional relationship "
 "from a nonspecific bulk-tissue immune-infiltration/cell-composition artifact. Cross-cohort "
 "synthesis of the primary endpoint used a random-effects meta-analysis implemented from first "
 "principles (Fisher z-transformation of Spearman rho, DerSimonian-Laird between-study variance "
 "estimator, τ²) (21), given the absence of a package dependency (metafor) in the analysis "
 "environment; results were verified to reproduce standard formulas. Heterogeneity was quantified "
 "with Cochran's Q and I². For the HCC-ICI cohorts, responder-versus-non-responder comparisons "
 "used the Wilcoxon rank-sum test, with effect size expressed as the scale-free rank-biserial "
 "correlation to allow pooling across cohorts measured on different expression "
 "scales/normalizations; the three cohorts were combined using the same random-effects "
 "meta-analysis framework. All statistical analyses were performed in R (version 4.5.2) using the "
 "DESeq2, GEOquery, org.Hs.eg.db, boot, and ggplot2 packages; in silico transcription-factor "
 "enrichment used the ChEA3 web API (22)."
)

doc.add_heading("2.5 In silico transcription-factor enrichment", level=2)
doc.add_paragraph(
 "To assess whether existing public ChIP-seq, co-expression, and literature-curated evidence "
 "supports HSF2 as a direct upstream regulator of the genes identified in this study, three gene "
 "sets — (i) the five genes found to be consistently negatively associated with HSF2 across "
 "all four fibrosis cohorts (PADI4, ELANE, MPO, CAMP, DEFA4), (ii) four genes found to have "
 "organ-divergent associations (LTF, NCF2, ITGAM, LCN2), and (iii) the pre-specified NETosis-core "
 "panel alone (PADI4, ELANE, MPO) — were submitted to ChEA3 (22), which ranks candidate "
 "upstream transcription factors against a query gene set across eight evidence libraries "
 "(including ENCODE ChIP-seq, ReMap ChIP-seq, Literature ChIP-seq, ARCHS4 co-expression, and GTEx "
 "co-expression). HSF2's rank and gene-overlap count were extracted from each library in which it "
 "was returned."
)

# =====================================================================
# 3. RESULTS
# =====================================================================
doc.add_heading("3. Results", level=1)

doc.add_heading("Table 1. Fibrosis cohort overview.", level=3)
add_table(
    ["Cohort", "Organ", "Role", "Platform", "n (total)", "n (fibrotic/diseased)", "Ref."],
    [
        ["GSE135251", "Liver", "Primary", "RNA-seq", "216", "206 (NAFLD, F0–F4)", "(14)"],
        ["GSE14323", "Liver", "Replication", "Affymetrix HG-U133A", "115", "96 (cirrhosis/cirrhosisHCC/HCC)", "(15)"],
        ["GSE47460", "Lung", "Primary", "Agilent GPL14550", "429 (213 used)", "122 (UIP/IPF)", "GEO accession"],
        ["GSE53845", "Lung", "Replication", "Agilent GPL6480 (2-color)", "48", "40 (IPF)", "(16)"],
    ],
    col_widths=[0.9, 0.55, 0.75, 1.3, 0.75, 1.4, 0.6],
)

doc.add_heading("3.1 HSF2 is negatively, not positively, associated with the NETosis-core "
                 "gene signature in fibrotic liver and lung tissue", level=2)
doc.add_paragraph(
 "Across all four independent fibrosis cohorts, the pre-specified primary endpoint — Spearman "
 "correlation between HSF2 and the NETosis-core score (PADI4/ELANE/MPO) within diseased/fibrotic "
 "tissue — was negative in direction, opposite to the original hypothesis (Table 2, Figure 1). The "
 "correlation reached nominal significance in three of four cohorts: GSE135251 (liver, n = 206 "
 "NAFLD samples), rho = −0.14 (95% CI −0.29 to 0.00), p = 0.038; GSE14323 (liver, n = 96 "
 "diseased samples), rho = −0.55 (95% CI −0.69 to −0.38), p = 6.2 × 10⁻⁹; "
 "and GSE47460 (lung, n = 122 UIP/IPF samples), rho = −0.39 (95% CI −0.54 to −0.22), "
 "p = 8.1 × 10⁻⁶. GSE53845 (lung, n = 40 IPF samples) showed a negative point estimate "
 "of smaller magnitude that did not reach significance (rho = −0.05, p = 0.76); given its "
 "markedly smaller sample size, this cohort was interpreted as underpowered rather than as "
 "evidence against the direction observed elsewhere. By the pre-specified replication rule, the "
 "liver primary/replication pair (GSE135251, GSE14323) met the criterion for replication (same "
 "sign, both p < 0.05); the lung pair did not formally meet the criterion because GSE53845 was not "
 "individually significant, though the direction was concordant. Full statistics for this endpoint, "
 "including bootstrap confidence intervals and leave-one-out sensitivity ranges for each cohort, "
 "are provided in Supplementary Table S1a; per-cohort scatter plots and boxplots underlying Table 2 "
 "and Figure 1 are provided in Supplementary Figures S1–S8."
)
doc.add_paragraph(
 "A random-effects meta-analysis of all four cohorts yielded a pooled rho of −0.31 (95% CI "
 "−0.51 to −0.08, p = 0.010), with substantial heterogeneity (I² = 83.5%) "
 "attributable largely to the smaller-magnitude GSE53845 estimate. Organ-restricted "
 "meta-analyses were directionally consistent but individually non-significant given only two "
 "studies each (liver-only pooled rho = −0.36, 95% CI −0.69 to 0.09, p = 0.11; lung-only "
 "pooled rho = −0.25, 95% CI −0.54 to 0.10, p = 0.15) (full meta-analysis output: Supplementary "
 "Table S1b)."
)

doc.add_heading("Table 2. Primary endpoint (HSF2 vs. NETosis-core score) results.", level=3)
add_table(
    ["Cohort", "n", "rho", "95% CI", "p"],
    [
        ["GSE135251 (liver, primary)", "206", "−0.145", "−0.288, −0.002", "0.038"],
        ["GSE14323 (liver, replication)", "96", "−0.550", "−0.689, −0.381", "6.2 × 10⁻⁹"],
        ["GSE47460 (lung, primary)", "122", "−0.392", "−0.538, −0.218", "8.1 × 10⁻⁶"],
        ["GSE53845 (lung, replication)", "40", "−0.050", "−0.350, 0.274", "0.760"],
        ["Pooled, random-effects (4 cohorts)", "464", "−0.308", "−0.508, −0.076", "0.010 (I²=83.5%)"],
    ],
    col_widths=[2.1, 0.5, 0.7, 1.3, 1.2],
)

add_figure(os.path.join(FIG, "FOREST_HSF2_vs_NETcore_all_cohorts.png"), width_in=6.0)
add_caption(
    "Forest plot of the pre-specified primary endpoint (Spearman rho, HSF2 vs. NETosis-core "
    "score) across all four fibrosis cohorts and pooled random-effects estimates. Error bars are "
    "bootstrap (cohort-level) or random-effects meta-analytic (pooled) 95% confidence intervals; "
    "point size is proportional to sample size. See also Supplementary Figures S1–S8 for the "
    "underlying per-cohort scatter plots and boxplots, and Supplementary Table S1a for full "
    "statistics.",
    bold_prefix="Figure 1. "
)

doc.add_heading("3.2 The negative association is not explained by generic immune-cell "
                 "infiltration", level=2)
doc.add_paragraph(
 "Because bulk-tissue expression reflects the mixture of resident and infiltrating cell types, an "
 "increase in overall immune infiltration accompanying advancing fibrosis could, in principle, "
 "dilute the relative contribution of non-immune, HSF2-expressing cells to the bulk signal — a "
 "trivial explanation unrelated to any specific regulatory relationship. This was tested directly "
 "using PTPRC (CD45) as a pan-leukocyte marker in the two liver cohorts. In both cohorts, HSF2 "
 "correlated positively, not negatively, with PTPRC (GSE135251: rho = 0.18, p = 0.010; GSE14323: "
 "rho = 0.45, p = 4.7 × 10⁻⁶) (Table 3), the opposite direction from its correlation "
 "with the NETosis-core score. This pattern — HSF2 rising with overall immune infiltration "
 "while falling specifically with NETosis-core gene expression — is inconsistent with a simple "
 "dilution artifact and instead suggests a signal specific to the NETosis gene program. Full "
 "statistics, including bootstrap confidence intervals and leave-one-out sensitivity ranges for "
 "all three pairwise correlations in both cohorts, are provided in Supplementary Table S1c."
)

doc.add_heading("Table 3. Specificity control: HSF2 vs. PTPRC (pan-leukocyte marker).", level=3)
add_table(
    ["Cohort", "HSF2 vs PTPRC, rho (p)", "HSF2 vs NET_core, rho (p)", "PTPRC vs NET_core, rho (p)"],
    [
        ["GSE135251 (liver)", "0.178 (0.010)", "−0.145 (0.038)", "0.252 (0.0003)"],
        ["GSE14323 (liver)", "0.448 (4.7 × 10⁻⁶)", "−0.550 (6.2 × 10⁻⁹)", "−0.224 (0.028)"],
    ],
    col_widths=[1.4, 1.7, 1.7, 1.7],
)

doc.add_heading("3.3 Individual-gene cross-cohort consistency identifies a reproducible core "
                 "and an organ-divergent set", level=2)
doc.add_paragraph(
 "Decomposing the NETosis-extended panel into individual genes and examining their direction of "
 "association with HSF2 across all four cohorts (Table 4; full per-cohort individual-gene results: "
 "Supplementary Tables S2; full 12-gene cross-cohort consistency table: Supplementary Table S2e) "
 "revealed two distinct patterns. Five genes — CAMP, MPO, DEFA4, and, with one small-magnitude "
 "exception each, PADI4 and ELANE — showed the same (negative) direction of association with HSF2 "
 "in all four cohorts, reaching FDR < 0.05 in two to three of four cohorts each. In contrast, four "
 "genes — LTF, NCF2, ITGAM, and LCN2 — showed sign reversal between liver and lung cohorts (e.g., "
 "LTF: negative in GSE135251 but positive and FDR-significant in both lung cohorts), a pattern "
 "that may reflect organ-specific regulation or a lesser degree of neutrophil-lineage specificity "
 "for these genes (LCN2 in particular is also broadly expressed by stressed hepatocytes and "
 "epithelial cells)."
)

doc.add_heading("Table 4. Cross-cohort consistency for the five most reproducible NETosis "
                 "panel genes (full 12-gene table: Supplementary Table S2e).", level=3)
add_table(
    ["Gene", "GSE135251 rho", "GSE14323 rho", "GSE47460 rho", "GSE53845 rho", "Cohorts FDR<0.05", "Direction"],
    [
        ["CAMP", "−0.056", "−0.459", "−0.457", "−0.529", "3/4", "Consistent (all −)"],
        ["MPO", "−0.025", "−0.277", "−0.320", "−0.412", "3/4", "Consistent (all −)"],
        ["DEFA4", "−0.090", "−0.337", "−0.328", "−0.097", "2/4", "Consistent (all −)"],
        ["PADI4", "−0.120", "−0.496", "−0.205", "+0.181", "2/4", "3 neg / 1 pos (n.s.)"],
        ["ELANE", "−0.192", "−0.422", "−0.348", "+0.075", "3/4", "3 neg / 1 pos (n.s.)"],
    ],
    col_widths=[0.55, 0.95, 0.95, 0.95, 0.95, 0.95, 1.2],
)

doc.add_heading("3.4 No existing in silico evidence supports HSF2 as a direct positive "
                 "regulator of these genes", level=2)
doc.add_paragraph(
 "Querying ChEA3 with the consistently associated gene set (PADI4, ELANE, MPO, CAMP, DEFA4), the "
 "organ-divergent gene set, and the NETosis-core panel alone, HSF2 did not rank among the top "
 "candidate regulators in any of the eight ChEA3 evidence libraries for any query (Table 5). Where "
 "HSF2 was returned at all, its rank fell in approximately the bottom 10% of all ranked "
 "transcription factors (e.g., Integrated mean rank 1444 of 1632 for the consistent gene set), "
 "with zero gene overlap (intersect = 0) and an odds ratio of 0 in every co-expression-based "
 "library. HSF2 was not returned at all by the ENCODE ChIP-seq, ReMap ChIP-seq, or Literature "
 "ChIP-seq libraries for any of the three queries. This absence of supporting evidence is "
 "consistent with two non-exclusive possibilities: (i) ChEA3's co-expression-based libraries are "
 "structurally unable to detect a negative/repressive regulatory relationship of the kind observed "
 "here, and (ii) this specific regulatory axis has not previously been characterized in the public "
 "ChIP-seq and functional genomics literature. Full results for all three queried gene sets across "
 "all eight ChEA3 libraries are provided in Supplementary Table S3."
)

doc.add_heading("Table 5. ChEA3 in silico enrichment: HSF2 rank summary (consistent gene set "
                 "query; full results for all 3 queries: Supplementary Table S3).", level=3)
add_table(
    ["Library", "HSF2 rank", "Total TFs ranked", "Gene overlap (intersect)"],
    [
        ["Integrated — mean rank", "1444", "1632", "—"],
        ["Integrated — top rank", "1383", "1632", "—"],
        ["GTEx co-expression", "1417", "1607", "0"],
        ["Enrichr Queries", "695", "1404", "0"],
        ["ARCHS4 co-expression", "1209", "1628", "0"],
        ["ENCODE ChIP-seq", "not returned", "118", "—"],
        ["ReMap ChIP-seq", "not returned", "297", "—"],
        ["Literature ChIP-seq", "not returned", "164", "—"],
    ],
    col_widths=[1.8, 1.3, 1.3, 1.6],
)

doc.add_heading("3.5 Exploratory analysis of HSF2 and response to HCC "
                 "immune-checkpoint-inhibitor therapy", level=2)
doc.add_paragraph(
 "Because HSF2 has previously been reported to correlate positively with immune-checkpoint gene "
 "expression in HCC (11), we performed an exploratory analysis of whether tumor HSF2 expression "
 "associates with clinical response to ICI therapy, using three independent cohorts with "
 "different regimens (Table 6, Figure 2). In GSE215011 (nivolumab monotherapy, n = 5 responders/5 "
 "non-responders), HSF2 did not differ significantly between groups (rank-biserial r = 0.44, "
 "p = 0.30). In GSE279750 (anti-PD-L1-based combination, post-treatment, n = 6/4), HSF2 was "
 "significantly lower in responders (r = −0.92, p = 0.025). In GSE235863 (anti-PD-1 plus "
 "lenvatinib, post-treatment, n = 11/4), HSF2 was also significantly lower in responders "
 "(r = −0.73, p = 0.043). A random-effects meta-analysis of all three cohorts was not "
 "significant (pooled r = −0.59, 95% CI −0.94 to 0.39, p = 0.22), with high heterogeneity "
 "(I² = 87.1%) driven by the opposite-direction, non-significant estimate from the "
 "monotherapy cohort. Given the small sample sizes (total n = 35 across three cohorts), differing "
 "regimens, and differing sampling timepoints (baseline versus post-treatment), this analysis is "
 "reported as an unresolved, hypothesis-generating lead rather than a confirmed finding. Full "
 "per-cohort sample-level data, group comparisons, and HSF2-vs-NET_core correlations are provided "
 "in Supplementary Tables S4; the underlying per-cohort scatter plots and boxplots are provided in "
 "Supplementary Figures S9–S15."
)

doc.add_heading("Table 6. HCC immune-checkpoint-inhibitor cohorts: HSF2, "
                 "responder vs. non-responder.", level=3)
add_table(
    ["Cohort", "Regimen", "n (R/NR)", "Rank-biserial r", "p"],
    [
        ["GSE215011", "anti-PD-1 monotherapy (nivolumab)", "5/5", "0.44", "0.30"],
        ["GSE279750", "anti-PD-L1 combination, post-tx", "6/4", "−0.92", "0.025"],
        ["GSE235863", "anti-PD-1 + lenvatinib combo, post-tx", "11/4", "−0.73", "0.043"],
        ["Pooled, random-effects", "3 cohorts combined", "35", "−0.59", "0.22 (I²=87.1%)"],
    ],
    col_widths=[1.1, 2.2, 0.8, 1.1, 1.2],
)

add_figure(os.path.join(FIG, "FOREST_hcc_ICI_HSF2_pooled.png"), width_in=6.0)
add_caption(
    "Forest plot of tumor HSF2 (responder vs. non-responder, rank-biserial correlation) across "
    "three independent HCC immune-checkpoint-inhibitor cohorts and the pooled random-effects "
    "estimate. Positive values indicate higher HSF2 in responders. See also Supplementary Figures "
    "S9–S15 for the underlying per-cohort scatter plots and boxplots, and Supplementary Tables "
    "S4 for full statistics.",
    bold_prefix="Figure 2. "
)

# =====================================================================
# 4. DISCUSSION
# =====================================================================
doc.add_heading("4. Discussion", level=1)

discussion_paragraphs = [
 "This study set out to test, using independent public transcriptomic data, a specific "
 "mechanistic prediction that followed from a wet-laboratory pilot observation: that "
 "toxicant-induced HSF2 upregulation in neutrophils drives NETosis-associated gene expression, "
 "thereby contributing to fibrosis and HCC immune evasion. Contrary to this prediction, HSF2 "
 "expression was consistently negatively, not positively, associated with a core NETosis gene "
 "signature across four independent cohorts spanning two organs and three expression-profiling "
 "platforms, with a statistically significant pooled effect and a pattern that survived a direct "
 "specificity control against generic immune infiltration. An in silico search of existing "
 "regulatory evidence found no support for HSF2 as a direct positive regulator of these genes in "
 "any of eight independent evidence libraries, consistent with this being either a repressive "
 "relationship (which co-expression-based enrichment tools are not designed to detect) or a "
 "genuinely uncharacterized regulatory axis, or both.",

 "Taken together, these results argue against the specific model in which HSF2 upregulation is a "
 "proximal driver of NETosis gene expression, and instead raise the alternative hypothesis that "
 "HSF2 may act as a negative regulator of, or be reciprocally regulated with, the NETosis "
 "transcriptional program in fibrotic tissue. This alternative is not inherently implausible: HSF2 "
 "is known to physically and functionally interact with HSF1 in a manner that can either enhance "
 "or repress HSF1-dependent transcription depending on cellular context (8), and context-dependent, "
 "even opposing, roles for HSF2 have been documented across other tissues and disease states "
 "(tumor-suppressive in prostate cancer, pro-malignant in cooperation with HSF1 in other cancers) "
 "(9, 10). The present findings do not contradict prior reports that HSF2 correlates positively "
 "with immune-checkpoint gene expression and confers poor prognosis in HCC via an "
 "EHMT2-FBP1-glycolysis axis intrinsic to tumor cells (11); rather, they suggest that any "
 "relationship between HSF2 and the neutrophil/NETosis compartment of the tumor or fibrotic "
 "microenvironment may run in a different, and possibly opposite, direction from HSF2's "
 "tumor-intrinsic oncogenic role — two mechanistically distinct axes that could plausibly "
 "coexist within the same disease.",

 "This study has several important limitations. First, and most fundamentally, all primary "
 "analyses used bulk tissue transcriptomes, which cannot establish which cell type(s) contribute "
 "the observed HSF2 or NETosis-gene signal, nor can they establish causal direction; the "
 "specificity control against PTPRC argues against a trivial infiltration-dilution explanation but "
 "cannot substitute for cell-type-resolved or single-cell data. Second, the analysis is "
 "cross-sectional and correlational; longitudinal or interventional data would be required to "
 "establish causality. Third, the smallest fibrosis cohort (GSE53845, n = 40) was underpowered and "
 "its non-significant result should not be over-interpreted as a true null. Fourth, the "
 "exploratory HCC-ICI analysis pooled cohorts that differed in treatment regimen (monotherapy "
 "versus combination therapies) and sampling timepoint (unclear versus post-treatment), "
 "introducing heterogeneity that likely explains why the pooled estimate did not reach "
 "significance despite two of three individual cohorts showing a significant effect; this analysis "
 "should be considered hypothesis-generating only. Fifth, the ChEA3 in silico analysis reflects "
 "the current state of public ChIP-seq and co-expression data, which is sparse for HSF2 (no "
 "ENCODE, ReMap, or literature-curated ChIP-seq data for HSF2 were available in myeloid or "
 "hepatic/pulmonary tissue contexts) and is not designed to detect repressive relationships; a "
 "null ChEA3 result is therefore evidence of absence of characterization, not evidence of absence "
 "of a real regulatory relationship.",

 "These findings directly inform the next stage of this research program. Rather than proceeding "
 "on the assumption that HSF2 upregulation will drive N2/NETosis marker expression in "
 "PHMG-treated neutrophils, the revised, pre-registered prediction motivated by the present "
 "analysis is that HSF2 induction by PHMG will not be accompanied by parallel increases in PADI4, "
 "ELANE, MPO, CAMP, and DEFA4 expression or in functional NET release, and may instead be "
 "associated with unchanged or reduced expression of these genes. This prediction, together with "
 "the standard N1 (iNOS, TNF-α, ICAM-1) and N2 (Arg-1, CD206, IL-10) polarization markers and a "
 "direct functional NETosis assay (e.g., extracellular DNA quantification), will be tested "
 "directly in isolated murine or human neutrophils across a PHMG dose-and-time matrix, providing a "
 "cell-type-resolved test of the hypothesis refined here.",
]
for para in discussion_paragraphs:
    doc.add_paragraph(para)

# =====================================================================
# 5. CONCLUSION
# =====================================================================
doc.add_heading("5. Conclusion", level=1)
doc.add_paragraph(
 "Across four independent public transcriptomic cohorts spanning liver and lung fibrosis, HSF2 "
 "expression was reproducibly and specifically negatively associated with a core NETosis gene "
 "signature, contradicting the original hypothesis that toxicant-induced HSF2 upregulation drives "
 "NETosis-associated gene expression. No existing in silico regulatory evidence supports a direct "
 "positive HSF2-NETosis gene relationship, and an exploratory analysis of HCC "
 "immune-checkpoint-inhibitor response cohorts provided an inconclusive, heterogeneity-limited "
 "signal warranting further investigation. These results motivate a revised, cell-type-resolved "
 "experimental design for the next phase of this research program."
)

# =====================================================================
# DATA AVAILABILITY
# =====================================================================
doc.add_heading("Data and Code Availability", level=1)
doc.add_paragraph(
 "All analysis code, the pre-specified statistical plan, intermediate and final result tables, "
 "and figures are available in a version-controlled repository (GitHub: bioinform25/project2). "
 "Raw data were obtained from the NCBI Gene Expression Omnibus under accession numbers GSE135251, "
 "GSE14323, GSE47460, GSE53845, GSE215011, GSE279750, GSE235863, and GSE140901, and are not "
 "redistributed in the repository owing to file size; download and reprocessing instructions are "
 "provided in the repository README."
)

doc.add_heading("Supplementary Materials", level=1)
doc.add_paragraph(
 "A companion Supplementary Materials document accompanies this manuscript, organized into five "
 "sections: S1, full fibrosis-cohort statistics (primary endpoint with bootstrap CI and "
 "leave-one-out sensitivity, meta-analysis, specificity control, and gene detection rates); S2, "
 "individual NETosis-panel gene results for each of the four fibrosis cohorts and the full "
 "12-gene cross-cohort consistency table; S3, the complete ChEA3 in silico enrichment output for "
 "all three queried gene sets; S4, full statistics for all three HCC immune-checkpoint-inhibitor "
 "cohorts; and S5, all cohort-level scatter and boxplot figures not reproduced in the main text "
 "(Supplementary Figures S1–S15)."
)

# =====================================================================
# REFERENCES
# =====================================================================
doc.add_page_break()
doc.add_heading("References", level=1)

references = [
 "Pelizzaro F, Farinati F, Trevisani F. Immune Checkpoint Inhibitors in Hepatocellular Carcinoma: "
 "Current Strategies and Biomarkers Predicting Response and/or Resistance. Biomedicines. "
 "2023;11(4):1020.",

 "Fridlender ZG, Sun J, Kim S, Kapoor V, Cheng G, Ling L, Worthen GS, Albelda SM. Polarization of "
 "tumor-associated neutrophil phenotype by TGF-β: “N1” versus “N2” TAN. Cancer Cell. "
 "2009;16(3):183-194.",

 "Ohms M, Möller S, Laskay T. An Attempt to Polarize Human Neutrophils Toward N1 and N2 "
 "Phenotypes in vitro. Front Immunol. 2020;11:532.",

 "Leshner M, Wang S, Lewis C, Zheng H, Chen XA, Santy L, Wang Y. PAD4 mediated histone "
 "hypercitrullination induces heterochromatin decondensation and chromatin unfolding to form "
 "neutrophil extracellular trap-like structures. Front Immunol. 2012;3:307.",

 "Jeong MH, Jeon MS, Kim GE, Kim HR. Polyhexamethylene Guanidine Phosphate Induces Apoptosis "
 "through Endoplasmic Reticulum Stress in Lung Epithelial Cells. Int J Mol Sci. 2021;22(3):1215.",

 "Kang MS, Kim SH, Yang MJ, Kim HY, Kim IH, Kang JW, Choi HS, Jin SW, Park EJ. Polyhexamethylene "
 "guanidine phosphate-induced necrosis may be linked to pulmonary fibrosis. Toxicol Lett. "
 "2022;362:1-16.",

 "Mathew A, Mathur SK, Jolly C, Fox SG, Kim S, Morimoto RI. Stress-Specific Activation and "
 "Repression of Heat Shock Factors 1 and 2. Mol Cell Biol. 2001;21(21):7163-7171.",

 "Sandqvist A, Björk JK, Åkerfelt M, Chitikova Z, Grichine A, Vourc'h C, Jolly C, Salminen TA, "
 "Nymalm Y, Sistonen L. Heterotrimerization of Heat-Shock Factors 1 and 2 Provides a "
 "Transcriptional Switch in Response to Distinct Stimuli. Mol Biol Cell. 2009;20(5):1340-1347.",

 "Björk JK, Åkerfelt M, Joutsen J, Puustinen MC, Cheng F, Sistonen L, Nees M. Heat-shock factor 2 "
 "is a suppressor of prostate cancer invasion. Oncogene. 2016;35(14):1770-1784.",

 "Smith RS, Takagishi SR, Amici DR, Metz K, Gayatri S, Alasady MJ, Wu Y, Brockway S, Taiberg SL, "
 "Khalatyan N, Taipale M, Santagata S, Whitesell L, Lindquist S, Savas JN, Mendillo ML. HSF2 "
 "cooperates with HSF1 to drive a transcriptional program critical for the malignant state. Sci "
 "Adv. 2022;8(11):eabj6526.",

 "Chen F, Fan Y, Liu X, Zhang J, Shang Y, Zhang B, Liu B, Hou J, Cao P, Tan K. Pan-Cancer "
 "Integrated Analysis of HSF2 Expression, Prognostic Value and Potential Implications for Cancer "
 "Immunity. Front Mol Biosci. 2022;8:789703.",

 "Barrett T, Wilhite SE, Ledoux P, Evangelista C, Kim IF, Tomashevsky M, Marshall KA, Phillippy "
 "KH, Sherman PM, Holko M, Yefanov A, Lee H, Zhang N, Robertson CL, Serova N, Davis S, Soboleva A. "
 "NCBI GEO: archive for functional genomics data sets--update. Nucleic Acids Res. "
 "2013;41(Database issue):D991-D995.",

 "Davis S, Meltzer PS. GEOquery: a bridge between the Gene Expression Omnibus (GEO) and "
 "BioConductor. Bioinformatics. 2007;23(14):1846-1847.",

 "Govaere O, Cockell S, Tiniakos D, Queen R, Younes R, Vacca M, Alexander L, Ravaioli F, Palmer J, "
 "Petta S, Boursier J, Rosso C, Johnson K, Wonders K, Day CP, Ekstedt M, Orešič M, Darlay R, "
 "Cordell HJ, Marra F, Vidal-Puig A, Bedossa P, Schattenberg JM, Clément K, Allison M, Bugianesi E, "
 "Ratziu V, Daly AK, Anstee QM. Transcriptomic profiling across the nonalcoholic fatty liver "
 "disease spectrum reveals gene signatures for steatohepatitis and fibrosis. Sci Transl Med. "
 "2020;12(572):eaba4448.",

 "Mas VR, Maluf DG, Archer KJ, Yanek K, Kong X, Kulik L, Freise CE, Olthoff KM, Ghobrial RM, "
 "McIver P, Fisher R. Genes involved in viral carcinogenesis and tumor initiation in hepatitis C "
 "virus-induced hepatocellular carcinoma. Mol Med. 2009;15(3-4):85-94.",

 "DePianto DJ, Chandriani S, Abbas AR, Jia G, N'Diaye EN, Caplazi P, Kauder SE, Biswas S, Karnik "
 "SK, Ha C, Modrusan Z, Matthay MA, Kukreja J, Collard HR, Egen JG, Wolters PJ, Arron JR. "
 "Heterogeneous gene expression signatures correspond to distinct lung pathologies and biomarkers "
 "of disease severity in idiopathic pulmonary fibrosis. Thorax. 2015;70(1):48-56.",

 "Liu C, Zhou C, Xia W, Zhou Y, Qiu Y, Weng J, Zhou Q, Chen W, Wang YN, Lee HH, Wang SC, Kuang M, "
 "Yu D, Ren N, Hung MC. Targeting ALK averts ribonuclease 1-induced immunosuppression and enhances "
 "antitumor immunity in hepatocellular carcinoma. Nat Commun. 2024;15:1009.",

 "Guo X, Nie H, Zhang W, Li J, Ge J, Xie B, Hu W, Zhu Y, Zhong N, Zhang X, Zhao X, Wang X, Sun Q, "
 "Wei K, Chen X, Ni L, Zhang T, Lu S, Zhang L, Dong C. Contrasting cytotoxic and regulatory T cell "
 "responses underlying distinct clinical outcomes to anti-PD-1 plus lenvatinib therapy in cancer. "
 "Cancer Cell. 2025;43(2):248-268.e9.",

 "Hsu CL, Ou DL, Bai LY, Chen CW, Lin L, Huang SF, Cheng AL, Jeng YM, Hsu C. Exploring Markers of "
 "Exhausted CD8 T Cells to Predict Response to Immune Checkpoint Inhibitor Therapy for "
 "Hepatocellular Carcinoma. Liver Cancer. 2021;10(4):346-359.",

 "Love MI, Huber W, Anders S. Moderated estimation of fold change and dispersion for RNA-seq data "
 "with DESeq2. Genome Biol. 2014;15:550.",

 "DerSimonian R, Laird N. Meta-analysis in clinical trials. Control Clin Trials. 1986;7(3):"
 "177-188.",

 "Keenan AB, Torre D, Lachmann A, Leong AK, Wojciechowicz ML, Utti V, Jagodnik KM, Kropiwnicki E, "
 "Wang Z, Ma'ayan A. ChEA3: transcription factor enrichment analysis by orthogonal omics "
 "integration. Nucleic Acids Res. 2019;47(W1):W212-W224.",

 "NCBI Gene Expression Omnibus, accession GSE279750. Anti-PD-L1-based combination immunotherapy "
 "response in hepatocellular carcinoma [dataset]. Accessed 2026. (No peer-reviewed publication was "
 "listed in the GEO record at the time of access.)",
]

for i, ref in enumerate(references, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    run = p.add_run(f"{i}. {ref}")
    run.font.size = Pt(10)

doc.save(OUT)
print(f"Manuscript build complete: {OUT}")
print(f"Total paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}")
